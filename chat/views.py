# chat/views.py - Updated to support frontend model selection
import os
import json
import requests
import base64
import io
import mimetypes

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from dotenv import load_dotenv
from django.shortcuts import render

load_dotenv()

# --------------------
# Model lists (single authoritative place)
# --------------------
VISION_MODELS = [
    "meta-llama/Llama-3.2-11B-Vision-Instruct",
    "Qwen/Qwen2.5-VL-7B-Instruct",
    "Qwen/Qwen2.5-VL-32B-Instruct",
    "CohereLabs/aya-vision-8b"
]

TEXT_MODELS = [
    "meta-llama/Llama-3.1-8B-Instruct",
    "deepseek-ai/DeepSeek-V3-0324",
    "Qwen/Qwen2.5-7B-Instruct"
]

# Optional libraries - import safely and handle missing packages gracefully
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import openpyxl
except Exception:
    openpyxl = None


def chat_page(request):
    """
    Render chat page and pass model lists so front-end can populate the dropdown.
    """
    context = {
        "vision_models": VISION_MODELS,
        "text_models": TEXT_MODELS,
    }
    return render(request, "chat/chat.html", context)


@csrf_exempt
def chat_view(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST allowed"}, status=405)

    try:
        # Determine model param from request (multipart or JSON)
        selected_model = None

        if request.content_type and request.content_type.startswith('multipart/form-data'):
            # File upload case
            message = request.POST.get('message', '').strip()
            selected_model = request.POST.get('model', '').strip() or None
            uploaded_file = request.FILES.get('file') or request.FILES.get('image')  # accept both 'file' and legacy 'image'

            if uploaded_file:
                return handle_uploaded_file(uploaded_file, message, selected_model)
            elif message:
                return handle_text_message(message, selected_model)
            else:
                return JsonResponse({"error": "No message or file provided"}, status=400)
        else:
            # JSON text message or base64 image case
            data = json.loads(request.body)
            message = data.get("message", "").strip()
            selected_model = data.get("model", None)

            # Check if there's base64 image data
            image_data = data.get("image")
            if image_data:
                return handle_base64_image_ocr(image_data, message, selected_model)
            elif message:
                return handle_text_message(message, selected_model)
            else:
                return JsonResponse({"error": "No message provided"}, status=400)

    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON in request"}, status=400)
    except Exception as e:
        print(f"Unexpected error: {e}")
        return JsonResponse({"error": "Internal server error"}, status=500)


def handle_uploaded_file(uploaded_file, message="", model=None):
    """
    Dispatch uploaded file to the appropriate extractor.
    Accepts images (delegates to Vision LLMs) and documents (pdf, docx, xlsx, csv, txt).
    model: string (model name) or None = auto
    """
    try:
        # Basic size limit (5MB) - adjust as needed
        max_bytes = 5 * 1024 * 1024
        if uploaded_file.size > max_bytes:
            return JsonResponse({"error": f"File too large. Max size is {max_bytes // (1024*1024)} MB."}, status=400)

        filename = uploaded_file.name
        content_type = uploaded_file.content_type or mimetypes.guess_type(filename)[0] or ""

        # Read bytes once
        file_bytes = uploaded_file.read()

        # If it's an image, use the vision LLM path
        if content_type.startswith("image/") or filename.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".webp", ".gif")):
            image_base64 = base64.b64encode(file_bytes).decode("utf-8")
            extracted_text = extract_text_with_vision_models(image_base64, message, model)
            if extracted_text:
                return JsonResponse({"response": extracted_text})
            else:
                return JsonResponse({"response": "I couldn't extract readable text from this image."})
        # Plain text
        elif content_type.startswith("text/") or filename.lower().endswith(".txt"):
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = file_bytes.decode("latin-1")
                except Exception:
                    text = None
            if text:
                return JsonResponse({"response": text})
            else:
                return JsonResponse({"response": "Could not decode text file."}, status=400)

        # CSV
        elif filename.lower().endswith(".csv") or content_type == "text/csv":
            try:
                decoded = file_bytes.decode("utf-8", errors="replace")
                lines = decoded.splitlines()
                preview = "\n".join(lines[:200])
                return JsonResponse({"response": preview})
            except Exception as e:
                print("CSV parsing error:", e)
                return JsonResponse({"response": "Failed to parse CSV file."}, status=500)

        # PDF
        elif filename.lower().endswith(".pdf") or content_type == "application/pdf":
            text = extract_text_from_pdf_bytes(file_bytes)
            if text and len(text.strip()) > 0:
                return JsonResponse({"response": text})
            else:
                return JsonResponse({"response": "No text could be extracted from the PDF."})

        # DOCX
        elif filename.lower().endswith(".docx") or content_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
            text = extract_text_from_docx_bytes(file_bytes)
            if text and len(text.strip()) > 0:
                return JsonResponse({"response": text})
            else:
                return JsonResponse({"response": "No text could be extracted from the DOCX file."})

        # XLSX
        elif filename.lower().endswith((".xlsx", ".xls")) or content_type in ("application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"):
            text = extract_text_from_xlsx_bytes(file_bytes)
            if text and len(text.strip()) > 0:
                return JsonResponse({"response": text})
            else:
                return JsonResponse({"response": "No text could be extracted from the spreadsheet."})

        else:
            # Unknown/unsupported type
            return JsonResponse({"response": f"Uploaded file type '{content_type or filename}' is not supported for automatic text extraction."}, status=400)

    except Exception as e:
        print(f"File handling error: {e}")
        return JsonResponse({"error": "Failed to process uploaded file"}, status=500)


def extract_text_from_pdf_bytes(file_bytes):
    """Try to extract text from PDF using PyPDF2 (if available)."""
    if not PyPDF2:
        print("PyPDF2 not installed")
        return "PDF text extraction unavailable (server missing PyPDF2). Install with: pip install PyPDF2"
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
                texts.append(page_text)
            except Exception:
                continue
        combined = "\n\n".join(texts).strip()
        return combined if len(combined) < 20000 else combined[:20000] + "\n\n[truncated]"
    except Exception as e:
        print("PDF extraction error:", e)
        return None


def extract_text_from_docx_bytes(file_bytes):
    """Extract plain text from DOCX using python-docx (if available)."""
    if not DocxDocument:
        print("python-docx not installed")
        return "DOCX extraction unavailable (server missing python-docx). Install with: pip install python-docx"
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                paragraphs.append(row_text)
        combined = "\n".join(paragraphs).strip()
        return combined if len(combined) < 20000 else combined[:20000] + "\n\n[truncated]"
    except Exception as e:
        print("DOCX extraction error:", e)
        return None


def extract_text_from_xlsx_bytes(file_bytes):
    """Extract text from XLSX using openpyxl (if available)."""
    if not openpyxl:
        print("openpyxl not installed")
        return "Spreadsheet extraction unavailable (server missing openpyxl). Install with: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(filename=io.BytesIO(file_bytes), data_only=True, read_only=True)
        texts = []
        for sheet in wb.worksheets:
            texts.append(f"--- Sheet: {sheet.title} ---")
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                row_count += 1
                row_str = ", ".join([str(cell) if cell is not None else "" for cell in row])
                texts.append(row_str)
                if row_count >= 200:
                    texts.append("[sheet truncated]")
                    break
        combined = "\n".join(texts).strip()
        return combined if len(combined) < 20000 else combined[:20000] + "\n\n[truncated]"
    except Exception as e:
        print("XLSX extraction error:", e)
        return None


# -------------------------
# Existing vision model OCR (adapted to accept specific model)
# -------------------------
def extract_text_with_vision_models(image_base64, user_message="", requested_model=None):
    """Extract text from image using Vision Language Models.
    If requested_model is provided and not 'auto', only try that model.
    """
    token = os.getenv("HUGGINGFACE_TOKEN")

    # Determine which vision models to try
    if requested_model and requested_model.lower() != "auto":
        models_to_try = [requested_model]
    else:
        models_to_try = VISION_MODELS

    # Create the prompt for OCR
    if user_message:
        prompt = f"Please extract all the text from this image and then answer this question: {user_message}"
    else:
        prompt = "Please extract and transcribe all the readable text from this image. Be as accurate as possible and maintain the original formatting when possible."

    for model_name in models_to_try:
        try:
            print(f"Trying Vision Model: {model_name}")

            api_url = "https://router.huggingface.co/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }

            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                    ]
                }
            ]

            payload = {
                "model": model_name,
                "messages": messages,
                "max_tokens": 500,
                "temperature": 0.1
            }

            response = requests.post(api_url, headers=headers, json=payload, timeout=45)
            print(f"Vision Model {model_name} status: {response.status_code}")

            if response.status_code == 200:
                result = response.json()
                print(f"Vision Model Success! Raw result: {result}")
                if "choices" in result and len(result["choices"]) > 0:
                    extracted_text = result["choices"][0]["message"]["content"]
                    if extracted_text and len(extracted_text.strip()) > 5:
                        return extracted_text.strip()

            elif response.status_code == 400:
                try:
                    error_info = response.json() if response.text else {}
                except Exception:
                    error_info = {}
                print(f"Vision Model {model_name} bad request: {error_info}")
                if "image_url" in str(error_info):
                    print(f"Trying simpler format for {model_name}")
                    return try_simple_vision_format(model_name, image_base64, prompt, token)
                continue

            elif response.status_code in (404, 503):
                print(f"Vision Model {model_name} returned {response.status_code}, trying next...")
                continue
            else:
                error_text = response.text[:200] if response.text else "No error details"
                print(f"Vision Model {model_name} error {response.status_code}: {error_text}")
                continue

        except requests.exceptions.Timeout:
            print(f"Timeout with vision model {model_name}, trying next...")
            continue
        except Exception as model_error:
            print(f"Error with vision model {model_name}: {model_error}")
            continue

    return None


def try_simple_vision_format(model_name, image_base64, prompt, token):
    """Try a simpler format for vision models that might not support the complex format"""
    try:
        api_url = "https://router.huggingface.co/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }

        messages = [
            {
                "role": "user",
                "content": f"{prompt}\n\n[Image provided as base64 data]",
                "image": image_base64
            }
        ]

        payload = {
            "model": model_name,
            "messages": messages,
            "max_tokens": 500,
            "temperature": 0.1
        }

        response = requests.post(api_url, headers=headers, json=payload, timeout=30)

        if response.status_code == 200:
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                extracted_text = result["choices"][0]["message"]["content"]
                if extracted_text and len(extracted_text.strip()) > 5:
                    return extracted_text.strip()

        print(f"Simple format also failed for {model_name}: {response.status_code}")
        return None

    except Exception as e:
        print(f"Simple format error for {model_name}: {e}")
        return None


def handle_base64_image_ocr(image_data, message="", model=None):
    """Handle OCR processing of base64 encoded image"""
    try:
        print("Processing base64 image data")
        if image_data.startswith('data:image'):
            image_data = image_data.split(',')[1]

        extracted_text = extract_text_with_vision_models(image_data, message, model)

        if extracted_text:
            return JsonResponse({"response": extracted_text})
        else:
            return JsonResponse({"response": "No readable text found in the image."})

    except Exception as e:
        print(f"Base64 OCR Error: {e}")
        return JsonResponse({"error": "Failed to process image data"}, status=500)


def handle_text_message(message, requested_model=None):
    """Handle regular text message (existing chat functionality), optionally using requested_model"""
    try:
        print(f"Processing text message: '{message}' with model={requested_model}")
        token = os.getenv("HUGGINGFACE_TOKEN")

        api_url = "https://router.huggingface.co/v1/chat/completions"
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

        # Choose models to try: either the single requested_model or the default list
        if requested_model and requested_model.lower() != "auto":
            models_to_try = [requested_model]
        else:
            models_to_try = TEXT_MODELS

        for model_name in models_to_try:
            try:
                payload = {
                    "model": model_name,
                    "messages": [{"role": "user", "content": message}],
                    "max_tokens": 150,
                    "temperature": 0.7,
                    "stream": False
                }

                response = requests.post(api_url, headers=headers, json=payload, timeout=30)

                if response.status_code == 200:
                    result = response.json()
                    if "choices" in result and len(result["choices"]) > 0:
                        answer = result["choices"][0]["message"]["content"]
                        return JsonResponse({"response": answer})

                elif response.status_code in [404, 503]:
                    continue
                else:
                    continue

            except Exception:
                continue

        return JsonResponse({"response": "I'm currently experiencing technical difficulties. Please try again in a moment."})

    except Exception as e:
        print(f"Text message error: {e}")
        return JsonResponse({"error": "Failed to process message"}, status=500)
